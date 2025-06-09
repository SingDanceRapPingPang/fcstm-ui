from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pyfcstm.model import State, CompositeState, PseudoState, Statechart

def center_text_in_cell(cell):
    """
    设置单元格中的文本居中显示
    
    Args:
        cell: 表格单元格对象
    """
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def export_statechart_to_word(statechart: Statechart, file_path: str) -> bool:
    """
    将状态机信息导出为Word文档
    
    Args:
        statechart: 状态机对象
        file_path: 导出文件路径
        
    Returns:
        bool: 是否导出成功
    """
    try:
        # 创建Word文档
        doc = Document()
        
        # 添加状态机信息表格
        table = doc.add_table(rows=6, cols=4)  # 创建6行4列的表格
        table.style = 'Table Grid'
        
        # 第一行：StateMachine
        cell = table.cell(0, 0)
        cell.merge(table.cell(0, 3))  # 合并第一行的所有列
        cell.text = "StateMachine"
        center_text_in_cell(cell)
        
        # 第二行：状态机名称
        table.cell(1, 0).text = "状态机名称"
        table.cell(1, 1).merge(table.cell(1, 3))  # 合并右侧单元格
        table.cell(1, 1).text = statechart.name
        center_text_in_cell(table.cell(1, 0))
        center_text_in_cell(table.cell(1, 1))
        
        # 第三行：状态机描述
        table.cell(2, 0).text = "状态机描述"
        table.cell(2, 1).merge(table.cell(2, 3))  # 合并右侧单元格
        table.cell(2, 1).text = '\n'.join(statechart.preamble)
        center_text_in_cell(table.cell(2, 0))
        center_text_in_cell(table.cell(2, 1))
        
        # 第四行：包含状态
        table.cell(3, 0).text = "包含状态"
        table.cell(3, 1).text = "状态名称"
        table.cell(3, 2).merge(table.cell(3, 3))  # 合并右侧单元格
        table.cell(3, 2).text = "状态类型"
        center_text_in_cell(table.cell(3, 0))
        center_text_in_cell(table.cell(3, 1))
        center_text_in_cell(table.cell(3, 2))
        
        # 添加状态信息
        row_idx = 4
        for state in statechart.states:
            if row_idx >= len(table.rows):
                table.add_row()
            table.cell(3, 0).merge(table.cell(row_idx, 0))
            table.cell(row_idx, 1).text = state.name
            table.cell(row_idx, 2).merge(table.cell(row_idx, 3))
            if isinstance(state, CompositeState):
                table.cell(row_idx, 2).text = "Composite"
            elif isinstance(state, PseudoState):
                table.cell(row_idx, 2).text = "Pseudo"
            else:
                table.cell(row_idx, 2).text = "Normal"
            center_text_in_cell(table.cell(row_idx, 1))
            center_text_in_cell(table.cell(row_idx, 2))
            row_idx += 1
        
        # 第五行：包含迁移
        if row_idx >= len(table.rows):
            table.add_row()
        transition_row_idx = row_idx
        table.cell(row_idx, 0).text = "包含迁移"
        table.cell(row_idx, 1).text = "起始状态"
        table.cell(row_idx, 2).text = "目标状态"
        table.cell(row_idx, 3).text = "激励事件"
        for i in range(4):
            center_text_in_cell(table.cell(row_idx, i))
        row_idx += 1
        
        # 添加迁移信息
        for transition in statechart.transitions:
            if row_idx >= len(table.rows):
                table.add_row()
            table.cell(transition_row_idx, 0).merge(table.cell(row_idx, 0))
            table.cell(row_idx, 1).text = transition.src_state.name
            table.cell(row_idx, 2).text = transition.dst_state.name
            table.cell(row_idx, 3).text = transition.event.name
            for i in range(1, 4):
                center_text_in_cell(table.cell(row_idx, i))
            row_idx += 1
        
        # 添加段落间距
        doc.add_paragraph()
        
        # 为每个状态添加详细信息
        for state in statechart.states:
            # 添加状态表格
            state_table = doc.add_table(rows=10, cols=3)  # 创建10行3列的表格
            state_table.style = 'Table Grid'
            
            # 第一行：State
            cell = state_table.cell(0, 0)
            cell.merge(state_table.cell(0, 2))  # 合并第一行的所有列
            cell.text = "State"
            center_text_in_cell(cell)
            
            # 第二行：状态名称
            state_table.cell(1, 0).text = "状态名称"
            state_table.cell(1, 1).merge(state_table.cell(1, 2))
            state_table.cell(1, 1).text = state.name
            center_text_in_cell(state_table.cell(1, 0))
            center_text_in_cell(state_table.cell(1, 1))
            
            # 第三行：状态描述
            state_table.cell(2, 0).text = "状态描述"
            state_table.cell(2, 1).merge(state_table.cell(2, 2))
            state_table.cell(2, 1).text = state.description if hasattr(state, 'description') else ""
            center_text_in_cell(state_table.cell(2, 0))
            center_text_in_cell(state_table.cell(2, 1))
            
            # 第四行：状态类型
            state_table.cell(3, 0).text = "状态类型"
            state_table.cell(3, 1).merge(state_table.cell(3, 2))
            if isinstance(state, CompositeState):
                state_table.cell(3, 1).text = "Composite"
            elif isinstance(state, PseudoState):
                state_table.cell(3, 1).text = "Pseudo"
            else:
                state_table.cell(3, 1).text = "Normal"
            center_text_in_cell(state_table.cell(3, 0))
            center_text_in_cell(state_table.cell(3, 1))
            
            # 第五行：时间锁
            state_table.cell(4, 0).text = "时间锁"
            state_table.cell(4, 1).text = "Min"
            state_table.cell(4, 2).text = state.min_time_lock if state.min_time_lock else ""
            state_table.cell(5, 0).text = ""
            state_table.cell(5, 1).text = "Max"
            state_table.cell(5, 2).text = state.max_time_lock if state.max_time_lock else ""
            state_table.cell(4, 0).merge(state_table.cell(5, 0))
            for i in range(3):
                center_text_in_cell(state_table.cell(4, i))
                center_text_in_cell(state_table.cell(5, i))
            
            # 第六行：Entry
            state_table.cell(6, 0).text = "Entry"
            state_table.cell(6, 1).merge(state_table.cell(6, 2))
            state_table.cell(6, 1).text = state.on_entry if state.on_entry else ""
            center_text_in_cell(state_table.cell(6, 0))
            center_text_in_cell(state_table.cell(6, 1))
            
            # 第七行：During
            state_table.cell(7, 0).text = "During"
            state_table.cell(7, 1).merge(state_table.cell(7, 2))
            state_table.cell(7, 1).text = state.during if hasattr(state, 'during') and state.during else ""
            center_text_in_cell(state_table.cell(7, 0))
            center_text_in_cell(state_table.cell(7, 1))
            
            # 第八行：Exit
            state_table.cell(8, 0).text = "Exit"
            state_table.cell(8, 1).merge(state_table.cell(8, 2))
            state_table.cell(8, 1).text = state.on_exit if state.on_exit else ""
            center_text_in_cell(state_table.cell(8, 0))
            center_text_in_cell(state_table.cell(8, 1))
            
            # 第九行：产生事件
            state_table.cell(9, 0).text = "产生事件"
            state_table.cell(9, 1).text = "事件名称"
            state_table.cell(9, 2).text = "事件产生条件"
            for i in range(3):
                center_text_in_cell(state_table.cell(9, i))
            
            # 添加事件信息
            row_idx = 10
            event_idx = row_idx - 1
            for transition in statechart.transitions:
                if transition.src_state_id == state.id:
                    if row_idx >= len(state_table.rows):
                        state_table.add_row()
                    state_table.cell(event_idx, 0).merge(state_table.cell(row_idx, 0))
                    state_table.cell(row_idx, 1).text = transition.event.name
                    state_table.cell(row_idx, 2).text = transition.event.guard if transition.event.guard else ""
                    for i in range(1, 3):
                        center_text_in_cell(state_table.cell(row_idx, i))
                    row_idx += 1
            if event_idx == row_idx - 1:
                state_table.add_row()
                state_table.cell(event_idx, 0).merge(state_table.cell(row_idx, 0))
                state_table.cell(row_idx, 1).text = ""
                state_table.cell(row_idx, 2).text = ""
                row_idx += 1
            transition_row_idx = row_idx
            # 第十行：迁移
            if row_idx >= len(state_table.rows):
                state_table.add_row()
            state_table.cell(row_idx, 0).text = "迁移"
            state_table.cell(row_idx, 1).text = "激励事件"
            state_table.cell(row_idx, 2).text = "迁移目标"
            for i in range(3):
                center_text_in_cell(state_table.cell(row_idx, i))
            row_idx += 1

            # 添加迁移信息
            for transition in statechart.transitions:
                if transition.src_state_id == state.id:
                    if row_idx >= len(state_table.rows):
                        state_table.add_row()
                    state_table.cell(transition_row_idx, 0).merge(state_table.cell(row_idx, 0))
                    state_table.cell(row_idx, 1).text = transition.event.name
                    state_table.cell(row_idx, 2).text = transition.dst_state.name
                    for i in range(1, 3):
                        center_text_in_cell(state_table.cell(row_idx, i))
                    row_idx += 1

            if transition_row_idx == row_idx - 1:
                state_table.add_row()
                state_table.cell(transition_row_idx, 0).merge(state_table.cell(row_idx, 0))
                state_table.cell(row_idx, 1).text = ""
                state_table.cell(row_idx, 2).text = ""
                row_idx += 1
            # 在每个状态表格后添加段落间距
            doc.add_paragraph()
        
        # 保存文档
        doc.save(file_path)
        return True
        
    except Exception as e:
        print(f"导出Word文档时发生错误：{str(e)}")
        return False
